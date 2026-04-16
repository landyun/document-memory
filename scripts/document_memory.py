#!/usr/bin/env python3
"""
Document Memory - 文档记忆与版本管理系统
支持多版本保存、变更追踪、差异对比
支持Markdown、纯文本和Word文档(.docx)格式
"""

import os
import json
import difflib
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any

# 尝试导入Word处理库
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DocumentMemory:
    """文档记忆管理器"""

    def __init__(self, project_name: str, project_path: str = "."):
        """
        初始化文档记忆项目

        Args:
            project_name: 项目名称
            project_path: 项目根路径
        """
        self.project_name = project_name
        self.project_path = Path(project_path)
        self.memory_root = self.project_path / ".doc-memory"
        self.project_id = self._sanitize_name(project_name)
        self.project_dir = self.memory_root / "projects" / self.project_id

        # 初始化目录结构
        self._init_directories()

        # 加载或创建项目元数据
        self.meta = self._load_or_create_meta()

    def _sanitize_name(self, name: str) -> str:
        """清理名称，用于文件路径"""
        return "".join(c if c.isalnum() or c in "-_" else "_" for c in name)

    def _init_directories(self):
        """初始化目录结构"""
        (self.project_dir / "versions").mkdir(parents=True, exist_ok=True)

    def _load_or_create_meta(self) -> Dict:
        """加载或创建项目元数据"""
        meta_path = self.project_dir / "meta.json"
        if meta_path.exists():
            with open(meta_path, "r", encoding="utf-8") as f:
                return json.load(f)

        # 创建新元数据
        meta = {
            "project_name": self.project_name,
            "project_id": self.project_id,
            "created_at": datetime.now().isoformat(),
            "current_version": 0,
            "description": ""
        }
        self._save_meta(meta)
        return meta

    def _save_meta(self, meta: Dict):
        """保存项目元数据"""
        with open(self.project_dir / "meta.json", "w", encoding="utf-8") as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)

    def _get_history_path(self) -> Path:
        """获取历史记录文件路径"""
        return self.project_dir / "history.json"

    def _load_history(self) -> List[Dict]:
        """加载版本历史"""
        history_path = self._get_history_path()
        if history_path.exists():
            with open(history_path, "r", encoding="utf-8") as f:
                return json.load(f)
        return []

    def _save_history(self, history: List[Dict]):
        """保存版本历史"""
        with open(self._get_history_path(), "w", encoding="utf-8") as f:
            json.dump(history, f, ensure_ascii=False, indent=2)

    def _get_version_path(self, version: int) -> Path:
        """获取版本文件路径"""
        return self.project_dir / "versions" / f"v{version}.md"

    def save_version(self, content: str, change_description: str = "",
                    author: str = "AI Assistant", metadata: Optional[Dict] = None) -> Dict:
        """
        保存新版本

        Args:
            content: 文档内容
            change_description: 变更描述
            author: 作者
            metadata: 附加元数据

        Returns:
            版本信息字典
        """
        # 递增版本号
        new_version = self.meta["current_version"] + 1

        # 保存内容
        version_path = self._get_version_path(new_version)
        with open(version_path, "w", encoding="utf-8") as f:
            f.write(content)

        # 创建版本记录
        version_info = {
            "version": new_version,
            "version_tag": f"v{new_version}",
            "timestamp": datetime.now().isoformat(),
            "change_description": change_description,
            "author": author,
            "metadata": metadata or {}
        }

        # 更新历史
        history = self._load_history()
        history.append(version_info)
        self._save_history(history)

        # 更新元数据
        self.meta["current_version"] = new_version
        self._save_meta(self.meta)

        return version_info

    def get_version(self, version: Optional[int] = None,
                   version_tag: Optional[str] = None) -> Optional[str]:
        """
        获取特定版本的内容

        Args:
            version: 版本号
            version_tag: 版本标签 (如 "v1")

        Returns:
            文档内容
        """
        if version_tag:
            version = int(version_tag.lstrip("v"))
        elif version is None:
            version = self.meta["current_version"]

        version_path = self._get_version_path(version)
        if version_path.exists():
            with open(version_path, "r", encoding="utf-8") as f:
                return f.read()
        return None

    def get_version_history(self) -> List[Dict]:
        """获取版本历史记录"""
        return self._load_history()

    def get_version_info(self, version: Optional[int] = None,
                        version_tag: Optional[str] = None) -> Optional[Dict]:
        """获取特定版本的信息"""
        if version_tag:
            version = int(version_tag.lstrip("v"))
        elif version is None:
            version = self.meta["current_version"]

        history = self._load_history()
        for info in history:
            if info["version"] == version:
                return info
        return None

    def compare_versions(self, version1: str, version2: str,
                        context_lines: int = 3) -> Dict[str, Any]:
        """
        对比两个版本

        Args:
            version1: 第一个版本标签 (如 "v1")
            version2: 第二个版本标签 (如 "v2")
            context_lines: 上下文行数

        Returns:
            差异分析结果
        """
        content1 = self.get_version(version_tag=version1)
        content2 = self.get_version(version_tag=version2)

        if content1 is None or content2 is None:
            return {"error": "无法找到一个或两个版本"}

        # 计算行级差异
        lines1 = content1.splitlines(keepends=True)
        lines2 = content2.splitlines(keepends=True)

        diff = difflib.unified_diff(
            lines1, lines2,
            fromfile=version1,
            tofile=version2,
            n=context_lines
        )

        diff_text = "".join(diff)

        # 统计变更
        added = 0
        removed = 0
        modified = 0

        for line in diff_text.splitlines():
            if line.startswith("+") and not line.startswith("+++"):
                added += 1
            elif line.startswith("-") and not line.startswith("---"):
                removed += 1

        # 生成摘要
        summary = {
            "added_lines": added,
            "removed_lines": removed,
            "total_changes": added + removed,
            "version1": version1,
            "version2": version2
        }

        # 分析章节变更
        section_changes = self._analyze_section_changes(content1, content2)

        return {
            "summary": summary,
            "diff_text": diff_text,
            "section_changes": section_changes,
            "content1": content1,
            "content2": content2
        }

    def _analyze_section_changes(self, content1: str, content2: str) -> List[Dict]:
        """分析章节级别的变更"""
        def extract_sections(content: str) -> Dict[str, str]:
            """提取Markdown章节"""
            sections = {}
            current_section = "前言"
            current_content = []

            for line in content.splitlines():
                if line.startswith("#"):
                    if current_section:
                        sections[current_section] = "\n".join(current_content)
                    current_section = line.lstrip("#").strip()
                    current_content = []
                else:
                    current_content.append(line)

            if current_section:
                sections[current_section] = "\n".join(current_content)

            return sections

        sections1 = extract_sections(content1)
        sections2 = extract_sections(content2)

        changes = []

        # 检查新增和修改的章节
        for title, content in sections2.items():
            if title not in sections1:
                changes.append({
                    "type": "added",
                    "section": title,
                    "content": content[:200] + "..." if len(content) > 200 else content
                })
            elif sections1[title] != content:
                changes.append({
                    "type": "modified",
                    "section": title,
                    "old": sections1[title][:100] + "..." if len(sections1[title]) > 100 else sections1[title],
                    "new": content[:100] + "..." if len(content) > 100 else content
                })

        # 检查删除的章节
        for title in sections1.keys():
            if title not in sections2:
                changes.append({
                    "type": "removed",
                    "section": title
                })

        return changes

    def generate_change_report(self, version1: str, version2: str) -> str:
        """
        生成变更报告

        Args:
            version1: 旧版本
            version2: 新版本

        Returns:
            Markdown格式的变更报告
        """
        diff_result = self.compare_versions(version1, version2)
        info1 = self.get_version_info(version_tag=version1)
        info2 = self.get_version_info(version_tag=version2)

        report = []
        report.append(f"# 变更报告: {version1} → {version2}\n")

        if info1 and info2:
            report.append(f"**{version1}**: {info1.get('change_description', '无描述')} ({info1.get('timestamp', '')[:19]})\n")
            report.append(f"**{version2}**: {info2.get('change_description', '无描述')} ({info2.get('timestamp', '')[:19]})\n")

        report.append("\n## 变更摘要\n")
        summary = diff_result["summary"]
        report.append(f"- 新增行: {summary['added_lines']}\n")
        report.append(f"- 删除行: {summary['removed_lines']}\n")
        report.append(f"- 总变更: {summary['total_changes']}\n")

        report.append("\n## 章节变更\n")
        for change in diff_result["section_changes"]:
            if change["type"] == "added":
                report.append(f"[+] **新增章节**: {change['section']}\n")
            elif change["type"] == "modified":
                report.append(f"[M] **修改章节**: {change['section']}\n")
            elif change["type"] == "removed":
                report.append(f"[-] **删除章节**: {change['section']}\n")

        report.append("\n## 详细差异\n")
        report.append("```diff\n")
        report.append(diff_result["diff_text"][:2000])
        if len(diff_result["diff_text"]) > 2000:
            report.append("\n... (更多内容已省略)\n")
        report.append("```\n")

        return "\n".join(report)

    def export_history(self, output_path: Optional[str] = None) -> str:
        """
        导出版本历史

        Args:
            output_path: 输出文件路径

        Returns:
            Markdown格式的历史记录
        """
        history = self.get_version_history()

        report = []
        report.append(f"# {self.project_name} - 版本历史\n")
        report.append(f"创建时间: {self.meta.get('created_at', '')[:19]}\n")
        report.append(f"当前版本: v{self.meta.get('current_version', 0)}\n")

        report.append("\n## 版本列表\n")

        for version_info in reversed(history):
            report.append(f"\n### {version_info['version_tag']}\n")
            report.append(f"- **时间**: {version_info['timestamp'][:19]}\n")
            report.append(f"- **作者**: {version_info['author']}\n")
            report.append(f"- **变更**: {version_info['change_description']}\n")

        content = "\n".join(report)

        if output_path:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)

        return content

    # ========== Word文档支持 ==========

    def import_from_word(self, docx_path: str, change_description: str = "",
                        author: str = "AI Assistant") -> Dict:
        """
        从Word文档(.docx)导入并保存为新版本

        Args:
            docx_path: Word文档路径
            change_description: 变更描述
            author: 作者

        Returns:
            版本信息字典
        """
        if not HAS_DOCX:
            raise ImportError("需要安装 python-docx 库: pip install python-docx")

        # 读取Word文档
        doc = Document(docx_path)
        content_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                # 检查是否是标题
                style_name = para.style.name
                if style_name.startswith('Heading'):
                    level = style_name[-1] if style_name[-1].isdigit() else '1'
                    content_parts.append(f"{'#' * int(level)} {para.text}")
                else:
                    content_parts.append(para.text)

        # 处理表格（可选，转换为简单文本）
        for table in doc.tables:
            content_parts.append("\n[表格]\n")
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                content_parts.append(row_text)

        content = "\n\n".join(content_parts)

        # 保存版本
        return self.save_version(content, change_description or f"从Word导入: {Path(docx_path).name}", author)

    def export_to_word(self, output_path: str, version: Optional[int] = None,
                      version_tag: Optional[str] = None) -> bool:
        """
        导出特定版本为Word文档(.docx)

        Args:
            output_path: 输出Word文档路径
            version: 版本号
            version_tag: 版本标签

        Returns:
            是否成功
        """
        if not HAS_DOCX:
            raise ImportError("需要安装 python-docx 库: pip install python-docx")

        content = self.get_version(version, version_tag)
        if content is None:
            return False

        doc = Document()

        # 设置默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

        # 解析Markdown并转换为Word格式
        lines = content.splitlines()
        i = 0

        while i < len(lines):
            line = lines[i].rstrip()

            if line.startswith('#'):
                # 处理标题
                level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()
                heading = doc.add_heading(title, level=min(level, 9))
                for run in heading.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')

            elif line.strip() == '':
                # 空行
                doc.add_paragraph()

            elif line.strip().startswith('[表格]'):
                # 跳过表格标记
                pass

            else:
                # 普通段落
                p = doc.add_paragraph(line)

            i += 1

        # 添加版本信息页脚
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        ver_info = self.get_version_info(version, version_tag)
        if ver_info:
            footer_para.text = f"版本: {ver_info['version_tag']} | 创建时间: {ver_info['timestamp'][:19]}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(output_path)
        return True

    def export_diff_to_word(self, output_path: str, version1: str, version2: str) -> bool:
        """
        导出版本差异对比为Word文档

        Args:
            output_path: 输出Word文档路径
            version1: 旧版本标签
            version2: 新版本标签

        Returns:
            是否成功
        """
        if not HAS_DOCX:
            raise ImportError("需要安装 python-docx 库: pip install python-docx")

        diff_result = self.compare_versions(version1, version2)
        info1 = self.get_version_info(version_tag=version1)
        info2 = self.get_version_info(version_tag=version2)

        doc = Document()

        # 设置默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

        # 标题
        title = doc.add_heading('版本差异对比报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 版本信息
        doc.add_heading('版本信息', level=1)

        info_table = doc.add_table(rows=2, cols=3)
        info_table.style = 'Light Grid Accent 1'
        info_table.rows[0].cells[0].text = '版本'
        info_table.rows[0].cells[1].text = '时间'
        info_table.rows[0].cells[2].text = '变更描述'

        if info1:
            info_table.rows[1].cells[0].text = info1['version_tag']
            info_table.rows[1].cells[1].text = info1['timestamp'][:19]
            info_table.rows[1].cells[2].text = info1['change_description']

        if info2:
            row_cells = info_table.add_row().cells
            row_cells[0].text = info2['version_tag']
            row_cells[1].text = info2['timestamp'][:19]
            row_cells[2].text = info2['change_description']

        # 变更摘要
        doc.add_heading('变更摘要', level=1)
        summary = diff_result['summary']
        p = doc.add_paragraph()
        p.add_run(f'新增行数: {summary["added_lines"]}\n').bold = True
        p.add_run(f'删除行数: {summary["removed_lines"]}\n').bold = True
        p.add_run(f'总变更数: {summary["total_changes"]}\n').bold = True

        # 章节变更
        doc.add_heading('章节变更', level=1)
        for change in diff_result['section_changes']:
            if change['type'] == 'added':
                p = doc.add_paragraph(f'[+] 新增章节: {change["section"]}', style='List Bullet')
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
            elif change['type'] == 'modified':
                p = doc.add_paragraph(f'[M] 修改章节: {change["section"]}', style='List Bullet')
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
            elif change['type'] == 'removed':
                p = doc.add_paragraph(f'[-] 删除章节: {change["section"]}', style='List Bullet')
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 0, 0)  # 红色

        # 详细差异
        doc.add_heading('详细差异', level=1)
        diff_para = doc.add_paragraph()
        diff_para.add_run('(完整差异请查看文本格式报告)').italic = True

        doc.save(output_path)
        return True


def main():
    """命令行接口"""
    import argparse

    parser = argparse.ArgumentParser(description="文档记忆与版本管理系统")
    subparsers = parser.add_subparsers(title="命令", dest="command")

    # 初始化项目
    init_parser = subparsers.add_parser("init", help="初始化新项目")
    init_parser.add_argument("--name", required=True, help="项目名称")
    init_parser.add_argument("--path", default=".", help="项目路径")

    # 保存版本
    save_parser = subparsers.add_parser("save", help="保存新版本")
    save_parser.add_argument("--name", required=True, help="项目名称")
    save_parser.add_argument("--file", required=True, help="文档文件")
    save_parser.add_argument("--message", required=True, help="变更描述")
    save_parser.add_argument("--author", default="AI Assistant", help="作者")

    # 列出版本
    list_parser = subparsers.add_parser("list", help="列出版本历史")
    list_parser.add_argument("--name", required=True, help="项目名称")

    # 对比版本
    diff_parser = subparsers.add_parser("diff", help="对比版本")
    diff_parser.add_argument("--name", required=True, help="项目名称")
    diff_parser.add_argument("version1", help="版本1 (如 v1)")
    diff_parser.add_argument("version2", help="版本2 (如 v2)")

    # 导出历史
    export_parser = subparsers.add_parser("export", help="导出版本历史")
    export_parser.add_argument("--name", required=True, help="项目名称")
    export_parser.add_argument("--output", help="输出文件")

    # 从Word导入
    import_word_parser = subparsers.add_parser("import-word", help="从Word文档(.docx)导入")
    import_word_parser.add_argument("--name", required=True, help="项目名称")
    import_word_parser.add_argument("--file", required=True, help="Word文档路径")
    import_word_parser.add_argument("--message", default="从Word导入", help="变更描述")
    import_word_parser.add_argument("--author", default="AI Assistant", help="作者")

    # 导出为Word
    export_word_parser = subparsers.add_parser("export-word", help="导出为Word文档(.docx)")
    export_word_parser.add_argument("--name", required=True, help="项目名称")
    export_word_parser.add_argument("--output", required=True, help="输出Word文档路径")
    export_word_parser.add_argument("--version", help="版本标签 (如 v1)，默认最新版本")

    # 导出差异为Word
    diff_word_parser = subparsers.add_parser("diff-word", help="导出版本差异为Word文档")
    diff_word_parser.add_argument("--name", required=True, help="项目名称")
    diff_word_parser.add_argument("--output", required=True, help="输出Word文档路径")
    diff_word_parser.add_argument("version1", help="旧版本 (如 v1)")
    diff_word_parser.add_argument("version2", help="新版本 (如 v2)")

    args = parser.parse_args()

    if args.command == "init":
        doc_memory = DocumentMemory(args.name, args.path)
        print(f"[OK] 项目 '{args.name}' 已初始化")

    elif args.command == "save":
        doc_memory = DocumentMemory(args.name)
        with open(args.file, "r", encoding="utf-8") as f:
            content = f.read()
        version_info = doc_memory.save_version(content, args.message, args.author)
        print(f"[OK] 已保存 {version_info['version_tag']}")

    elif args.command == "list":
        doc_memory = DocumentMemory(args.name)
        history = doc_memory.get_version_history()
        for v in history:
            print(f"{v['version_tag']:4} | {v['timestamp'][:19]} | {v['author']:20} | {v['change_description']}")

    elif args.command == "diff":
        doc_memory = DocumentMemory(args.name)
        report = doc_memory.generate_change_report(args.version1, args.version2)
        print(report)

    elif args.command == "export":
        doc_memory = DocumentMemory(args.name)
        content = doc_memory.export_history(args.output)
        if not args.output:
            print(content)
        else:
            print(f"[OK] 已导出到 {args.output}")

    elif args.command == "import-word":
        if not HAS_DOCX:
            print("[ERROR] 需要安装 python-docx: pip install python-docx")
            exit(1)
        doc_memory = DocumentMemory(args.name)
        version_info = doc_memory.import_from_word(args.file, args.message, args.author)
        print(f"[OK] 已从Word导入并保存为 {version_info['version_tag']}")

    elif args.command == "export-word":
        if not HAS_DOCX:
            print("[ERROR] 需要安装 python-docx: pip install python-docx")
            exit(1)
        doc_memory = DocumentMemory(args.name)
        if args.version:
            success = doc_memory.export_to_word(args.output, version_tag=args.version)
        else:
            success = doc_memory.export_to_word(args.output)
        if success:
            print(f"[OK] 已导出Word文档: {args.output}")
        else:
            print("[ERROR] 导出失败")

    elif args.command == "diff-word":
        if not HAS_DOCX:
            print("[ERROR] 需要安装 python-docx: pip install python-docx")
            exit(1)
        doc_memory = DocumentMemory(args.name)
        success = doc_memory.export_diff_to_word(args.output, args.version1, args.version2)
        if success:
            print(f"[OK] 已导出差异对比Word文档: {args.output}")
        else:
            print("[ERROR] 导出失败")


if __name__ == "__main__":
    main()
